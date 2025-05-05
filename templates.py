from fastapi import APIRouter, HTTPException, Body
import os
import json
from docx import Document
import uuid
from typing import Dict

router = APIRouter()

# Path to internal templates folder (updated from external folder)
TEMPLATES_FOLDER = os.path.join(os.path.dirname(__file__), "templates")

# Path to save generated documents
GENERATED_DOCS_FOLDER = os.path.join(os.path.dirname(__file__), "temp_audio")
os.makedirs(GENERATED_DOCS_FOLDER, exist_ok=True)

@router.get("/list")
def list_templates():
    """
    List available templates and their blocks.
    Assumes each template is a JSON file with a structure like:
    {
        "template_name": "Template 1",
        "blocks": [
            {"id": "block1", "name": "Introduction"},
            {"id": "block2", "name": "Body"},
            ...
        ]
    }
    """
    try:
        templates = []
        for filename in os.listdir(TEMPLATES_FOLDER):
            if filename.endswith(".json"):
                filepath = os.path.join(TEMPLATES_FOLDER, filename)
                with open(filepath, "r", encoding="utf-8") as f:
                    template_data = json.load(f)
                    templates.append({
                        "filename": filename,
                        "template_name": template_data.get("template_name", filename),
                        "blocks": template_data.get("blocks", [])
                    })
        # Return as JSON serializable object
        return {"templates": templates}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading templates: {str(e)}")

@router.post("/format")
def format_transcription(
    template_filename: str = Body(...),
    blocks_text: Dict[str, str] = Body(...)
):
    """
    Format the transcription text according to the selected template and blocks.
    Generates a .docx document with the transcription text inserted in the respective blocks.
    Returns a download URL for the generated document.
    """
    try:
        filepath = os.path.join(TEMPLATES_FOLDER, template_filename)
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="Template not found")

        with open(filepath, "r", encoding="utf-8") as f:
            template_data = json.load(f)

        # Create a new Word document
        doc = Document()

        # Add template name as title
        doc.add_heading(template_data.get("template_name", "Template"), level=1)

        # For each block in template, add block name and corresponding text
        for block in template_data.get("blocks", []):
            block_id = block.get("id")
            block_name = block.get("name", "Block")
            text = blocks_text.get(block_id, "")

            doc.add_heading(block_name, level=2)
            doc.add_paragraph(text)

        # Generate unique filename for the docx
        doc_filename = f"{uuid.uuid4()}.docx"
        doc_filepath = os.path.join(GENERATED_DOCS_FOLDER, doc_filename)

        # Save the document
        doc.save(doc_filepath)

        # Return download URL
        download_url = f"/temp_audio/{doc_filename}"

        return {
            "download_url": download_url
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error formatting transcription: {str(e)}")
        
from fastapi import APIRouter, HTTPException, Body
import os
import json
from docx import Document
import uuid
from typing import Dict

router = APIRouter()

# Path to external templates folder
TEMPLATES_FOLDER = r"C:\Users\aasav\Downloads\voice2text templates"

# Path to save generated documents
GENERATED_DOCS_FOLDER = os.path.join(os.path.dirname(__file__), "temp_audio")
os.makedirs(GENERATED_DOCS_FOLDER, exist_ok=True)

@router.get("/list")
def list_templates():
    """
    List available templates and their blocks.
    Assumes each template is a JSON file with a structure like:
    {
        "template_name": "Template 1",
        "blocks": [
            {"id": "block1", "name": "Introduction"},
            {"id": "block2", "name": "Body"},
            ...
        ]
    }
    """
    try:
        templates = []
        for filename in os.listdir(TEMPLATES_FOLDER):
            if filename.endswith(".json"):
                filepath = os.path.join(TEMPLATES_FOLDER, filename)
                with open(filepath, "r", encoding="utf-8") as f:
                    template_data = json.load(f)
                    templates.append({
                        "filename": filename,
                        "template_name": template_data.get("template_name", filename),
                        "blocks": template_data.get("blocks", [])
                    })
        # Return as JSON serializable object
        return {"templates": templates}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading templates: {str(e)}")

@router.post("/format")
def format_transcription(
    template_filename: str = Body(...),
    blocks_text: Dict[str, str] = Body(...)
):
    """
    Format the transcription text according to the selected template and blocks.
    Generates a .docx document with the transcription text inserted in the respective blocks.
    Returns a download URL for the generated document.
    """
    try:
        filepath = os.path.join(TEMPLATES_FOLDER, template_filename)
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="Template not found")

        with open(filepath, "r", encoding="utf-8") as f:
            template_data = json.load(f)

        # Create a new Word document
        doc = Document()

        # Add template name as title
        doc.add_heading(template_data.get("template_name", "Template"), level=1)

        # For each block in template, add block name and corresponding text
        for block in template_data.get("blocks", []):
            block_id = block.get("id")
            block_name = block.get("name", "Block")
            text = blocks_text.get(block_id, "")

            doc.add_heading(block_name, level=2)
            doc.add_paragraph(text)

        # Generate unique filename for the docx
        doc_filename = f"{uuid.uuid4()}.docx"
        doc_filepath = os.path.join(GENERATED_DOCS_FOLDER, doc_filename)

        # Save the document
        doc.save(doc_filepath)

        # Return download URL
        download_url = f"/temp_audio/{doc_filename}"

        return {
            "download_url": download_url
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error formatting transcription: {str(e)}")
